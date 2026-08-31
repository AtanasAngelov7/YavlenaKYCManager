from __future__ import annotations

import hashlib
import json
import os
import zipfile
from datetime import date, datetime, timezone
from pathlib import Path

import pytest
from pydantic import ValidationError

from bulgarian_numbers import bulgarian_integer_words
from contracts import ContractGenerationError, TEMPLATE_BY_ROLE, generate_contract
from models import (
    AgentDetails,
    BinaryChoice,
    ContactDetails,
    ContractInput,
    ContractOptions,
    ContractRole,
    PersonalDocument,
    PropertyDetailsSource,
    PropertyDocument,
    PropertyDocumentType,
    PropertyExtractionMethod,
    PropertyExtractionResult,
    personal_document_fingerprint,
)
from storage import write_json


PROPERTY_SOURCE_BYTES = b"authorized synthetic property document"
PROPERTY_SOURCE_SHA256 = hashlib.sha256(PROPERTY_SOURCE_BYTES).hexdigest()


def _contract_input(case_id: str, role: ContractRole) -> ContractInput:
    client = PersonalDocument(
        first_name="ИВАН",
        middle_name="ПЕТРОВ",
        last_name="ИВАНОВ",
        personal_number="6101057509",
        document_number="123456789",
    )
    extraction = _property_extraction(case_id, client)
    extraction_text = json.dumps(
        extraction.model_dump(mode="json"),
        ensure_ascii=False,
        indent=2,
        default=str,
    )
    extraction_bytes = extraction_text.replace("\n", os.linesep).encode("utf-8")
    seller_values = (
        {
            "property_details_source": PropertyDetailsSource.NOTARY_DOCUMENT,
            "property_description": "АПАРТАМЕНТ № 6 в гр. София, площ 160.20 кв. м.",
            "property_document_filename": "property-document.pdf",
            "property_document_sha256": PROPERTY_SOURCE_SHA256,
            "property_extraction_record_sha256": hashlib.sha256(extraction_bytes).hexdigest(),
            "property_extraction_method": PropertyExtractionMethod.STANDARD,
            "property_document_type": PropertyDocumentType.OWNERSHIP_NOTARIAL_ACT,
            "exclusive_term": "6 месеца",
            "offer_price_eur": "250 000",
            "offer_price_eur_words": "двеста и петдесет хиляди",
        }
        if role is ContractRole.SELLER
        else {}
    )
    return ContractInput(
        case_id=case_id,
        role=role,
        client=client,
        client_contact=ContactDetails(phone="+359881234567", email="client@example.test"),
        agent=AgentDetails(
            name="ТЕСТОВ АГЕНТ",
            phone="+359887654321",
            email="agent@example.test",
        ),
        options=ContractOptions(
            contract_date=date(2026, 8, 28),
            privacy_paper_choice=BinaryChoice.NO,
            privacy_email_choice=BinaryChoice.YES,
            privacy_email="client@example.test",
            marketing_choice=BinaryChoice.NO,
            **seller_values,
        ),
        approved_by_operator=True,
        approved_at=datetime(2026, 8, 28, 12, tzinfo=timezone.utc),
    )


def _word_xml_text(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return "".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )


def _property_extraction(case_id: str, client: PersonalDocument) -> PropertyExtractionResult:
    return PropertyExtractionResult(
        case_id=case_id,
        document=PropertyDocument(
            document_type=PropertyDocumentType.OWNERSHIP_NOTARIAL_ACT,
            property_description="АПАРТАМЕНТ № 6 в гр. София, площ 160.20 кв. м.",
        ),
        seller_identity_fingerprint=personal_document_fingerprint(client),
        source_filename="property-document.pdf",
        source_sha256=PROPERTY_SOURCE_SHA256,
        extraction_method=PropertyExtractionMethod.STANDARD,
    )


def _write_active_property_source(case_root: Path) -> None:
    source = case_root / "original" / "property-document.pdf"
    source.parent.mkdir(parents=True, exist_ok=True)
    source.write_bytes(PROPERTY_SOURCE_BYTES)
    client = _contract_input(case_root.name, ContractRole.SELLER).client
    write_json(
        case_root / "property_extracted.json",
        _property_extraction(case_root.name, client).model_dump(mode="json"),
    )


def _without_notary_provenance(options: ContractOptions) -> dict[str, object]:
    return {
        **options.model_dump(mode="python"),
        "property_document_filename": "",
        "property_document_sha256": "",
        "property_extraction_record_sha256": "",
        "property_extraction_method": None,
        "property_document_type": None,
        "property_ai_model": "",
        "property_ai_prompt_version": "",
        "property_ai_input_sha256": "",
        "property_ai_response_sha256": "",
        "property_external_processing_authorized_at": None,
    }


def test_generate_buyer_contract_bundle(tmp_path: Path) -> None:
    case_id = "2026-08-28_120000_abc123"
    case_root = tmp_path / case_id

    generated = generate_contract(_contract_input(case_id, ContractRole.BUYER), case_root)

    assert generated.document_path.is_file()
    assert generated.input_path.is_file()
    assert generated.manifest_path.is_file()
    rendered = _word_xml_text(generated.document_path)
    assert "ИВАН" in rendered
    assert "6101057509" in rendered
    assert "123456789" in rendered
    assert "client@example.test" in rendered
    assert "ДА ☒  НЕ ☐" in rendered
    assert "{{" not in rendered
    assert "град" in rendered  # Buyer search criteria remain manually fillable.

    stored_input = json.loads(generated.input_path.read_text(encoding="utf-8"))
    stored_manifest = json.loads(generated.manifest_path.read_text(encoding="utf-8"))
    assert stored_input["approved_by_operator"] is True
    assert stored_manifest["output_filename"] == generated.document_path.name
    assert stored_manifest["input_sha256"] == hashlib.sha256(
        generated.input_path.read_bytes()
    ).hexdigest()
    assert stored_manifest["output_sha256"] == hashlib.sha256(
        generated.document_path.read_bytes()
    ).hexdigest()


def test_unreviewed_poc_generation_is_recorded_without_fabricated_approval(
    tmp_path: Path,
) -> None:
    approved = _contract_input("2026-08-28_120050_poc001", ContractRole.BUYER)
    unreviewed = ContractInput.model_validate(
        {
            **approved.model_dump(mode="python"),
            "approved_by_operator": False,
            "approved_at": None,
            "warning_codes": ["synthetic_poc_warning"],
            "warnings_acknowledged": False,
        }
    )

    generated = generate_contract(unreviewed, tmp_path / unreviewed.case_id)
    stored_input = json.loads(generated.input_path.read_text(encoding="utf-8"))

    assert stored_input["approved_by_operator"] is False
    assert stored_input["approved_at"] is None
    assert stored_input["warnings_acknowledged"] is False


def test_generate_seller_contract_uses_approved_property_text(tmp_path: Path) -> None:
    case_id = "2026-08-28_120100_def456"
    case_root = tmp_path / case_id
    _write_active_property_source(case_root)

    generated = generate_contract(_contract_input(case_id, ContractRole.SELLER), case_root)

    rendered = _word_xml_text(generated.document_path)
    assert "АПАРТАМЕНТ № 6" in rendered
    assert "250 000" in rendered
    assert "Физическо лице" in rendered
    assert "Няколко физически лица" not in rendered

    stored_input = json.loads(generated.input_path.read_text(encoding="utf-8"))
    assert stored_input["options"]["property_details_source"] == "notary_document"
    assert stored_input["options"]["property_document_filename"] == "property-document.pdf"
    assert stored_input["options"]["property_document_sha256"] == PROPERTY_SOURCE_SHA256
    assert stored_input["options"]["property_extraction_method"] == "standard"
    assert stored_input["options"]["property_document_type"] == "ownership_notarial_act"
    assert stored_input["options"]["property_extraction_record_sha256"] == hashlib.sha256(
        (case_root / "property_extracted.json").read_bytes()
    ).hexdigest()


def test_seller_generation_requires_the_exact_active_property_source(tmp_path: Path) -> None:
    contract_input = _contract_input("2026-08-28_120150_source1", ContractRole.SELLER)
    case_root = tmp_path / contract_input.case_id

    with pytest.raises(ContractGenerationError, match="no longer active"):
        generate_contract(contract_input, case_root)

    source = case_root / "original" / "property-document.pdf"
    source.parent.mkdir(parents=True)
    source.write_bytes(b"different property document")
    with pytest.raises(ContractGenerationError, match="changed after extraction"):
        generate_contract(contract_input, case_root)


def test_seller_generation_rejects_a_changed_property_extraction_record(tmp_path: Path) -> None:
    case_id = "2026-08-28_120155_record1"
    case_root = tmp_path / case_id
    contract_input = _contract_input(case_id, ContractRole.SELLER)
    _write_active_property_source(case_root)
    record_path = case_root / "property_extracted.json"
    record = json.loads(record_path.read_text(encoding="utf-8"))
    record["document"]["property_description"] = "changed after review"
    write_json(record_path, record)

    with pytest.raises(ContractGenerationError, match="extraction changed after contract review"):
        generate_contract(contract_input, case_root)


def test_generation_never_overwrites_an_earlier_draft(tmp_path: Path) -> None:
    case_id = "2026-08-28_120200_ghi789"
    case_root = tmp_path / case_id
    contract_input = _contract_input(case_id, ContractRole.BUYER)

    first = generate_contract(contract_input, case_root)
    second = generate_contract(contract_input, case_root)

    assert first.document_path != second.document_path
    assert first.document_path.is_file()
    assert second.document_path.is_file()


def test_seller_requires_all_role_specific_values() -> None:
    buyer = _contract_input("2026-08-28_120300_jkl012", ContractRole.BUYER)

    with pytest.raises(ValidationError, match="Missing seller values"):
        ContractInput.model_validate(
            {
                **buyer.model_dump(mode="python"),
                "role": ContractRole.SELLER,
            }
        )


def test_notary_document_source_requires_an_auditable_file() -> None:
    seller = _contract_input("2026-08-28_120310_jkl013", ContractRole.SELLER)

    with pytest.raises(ValidationError, match="processed notary document"):
        ContractInput.model_validate(
            {
                **seller.model_dump(mode="python"),
                "options": {
                    **_without_notary_provenance(seller.options),
                },
            }
        )


def test_openai_property_source_requires_complete_ai_provenance() -> None:
    seller = _contract_input("2026-08-28_120315_ai0001", ContractRole.SELLER)

    with pytest.raises(ValidationError, match="complete AI provenance"):
        ContractInput.model_validate(
            {
                **seller.model_dump(mode="python"),
                "options": {
                    **seller.options.model_dump(mode="python"),
                    "property_extraction_method": PropertyExtractionMethod.OPENAI,
                },
            }
        )


def test_manual_property_source_requires_warning_and_acknowledgement() -> None:
    seller = _contract_input("2026-08-28_120320_jkl014", ContractRole.SELLER)
    manual_options = {
        **_without_notary_provenance(seller.options),
        "property_details_source": PropertyDetailsSource.MANUAL,
    }

    with pytest.raises(ValidationError, match="explicit review warning"):
        ContractInput.model_validate(
            {
                **seller.model_dump(mode="python"),
                "options": manual_options,
            }
        )

    with pytest.raises(ValidationError, match="warning must be acknowledged"):
        ContractInput.model_validate(
            {
                **seller.model_dump(mode="python"),
                "options": manual_options,
                "warning_codes": ["manual_property_details"],
                "warnings_acknowledged": False,
            }
        )


def test_manual_property_source_can_generate_after_explicit_review(tmp_path: Path) -> None:
    seller = _contract_input("2026-08-28_120330_jkl015", ContractRole.SELLER)
    manual = ContractInput.model_validate(
        {
            **seller.model_dump(mode="python"),
            "options": {
                **_without_notary_provenance(seller.options),
                "property_details_source": PropertyDetailsSource.MANUAL,
            },
            "warning_codes": ["manual_property_details"],
            "warnings_acknowledged": True,
        }
    )

    generated = generate_contract(manual, tmp_path / manual.case_id)

    assert "АПАРТАМЕНТ № 6" in _word_xml_text(generated.document_path)
    stored_input = json.loads(generated.input_path.read_text(encoding="utf-8"))
    assert stored_input["options"]["property_details_source"] == "manual"
    assert stored_input["warning_codes"] == ["manual_property_details"]
    assert stored_input["warnings_acknowledged"] is True


def test_manual_property_source_can_generate_as_truthfully_unreviewed_poc(
    tmp_path: Path,
) -> None:
    seller = _contract_input("2026-08-28_120335_jkl016", ContractRole.SELLER)
    unreviewed = ContractInput.model_validate(
        {
            **seller.model_dump(mode="python"),
            "options": {
                **_without_notary_provenance(seller.options),
                "property_details_source": PropertyDetailsSource.MANUAL,
            },
            "approved_by_operator": False,
            "approved_at": None,
            "warning_codes": ["manual_property_details"],
            "warnings_acknowledged": False,
        }
    )

    generated = generate_contract(unreviewed, tmp_path / unreviewed.case_id)
    stored_input = json.loads(generated.input_path.read_text(encoding="utf-8"))

    assert stored_input["approved_by_operator"] is False
    assert stored_input["warning_codes"] == ["manual_property_details"]
    assert stored_input["warnings_acknowledged"] is False


def test_critical_warnings_require_acknowledgement() -> None:
    buyer = _contract_input("2026-08-28_120400_mno345", ContractRole.BUYER)

    with pytest.raises(ValidationError, match="warning must be acknowledged"):
        ContractInput.model_validate(
            {
                **buyer.model_dump(mode="python"),
                "warning_codes": ["property_document_mismatch"],
                "warnings_acknowledged": False,
            }
        )


def test_contract_boundary_rejects_nonempty_but_invalid_identity_values() -> None:
    buyer = _contract_input("2026-08-28_120405_invalid", ContractRole.BUYER)

    with pytest.raises(ValidationError, match="Invalid client identity values"):
        ContractInput.model_validate(
            {
                **buyer.model_dump(mode="python"),
                "client": {
                    **buyer.client.model_dump(mode="python"),
                    "personal_number": "0000000000",
                    "document_number": "INVALID",
                },
            }
        )


def test_declined_privacy_email_is_not_retained_or_rendered() -> None:
    buyer = _contract_input("2026-08-28_120410_privacy", ContractRole.BUYER)
    declined = ContractInput.model_validate(
        {
            **buyer.model_dump(mode="python"),
            "options": {
                **buyer.options.model_dump(mode="python"),
                "privacy_email_choice": BinaryChoice.NO,
                "privacy_email": "unused and invalid",
            },
        }
    )

    assert declined.options.privacy_email == ""


def test_seller_price_rejects_non_numeric_values() -> None:
    with pytest.raises(ValidationError, match="positive whole-EUR price"):
        ContractOptions(
            contract_date=date(2026, 8, 28),
            privacy_paper_choice=BinaryChoice.NO,
            privacy_email_choice=BinaryChoice.NO,
            marketing_choice=BinaryChoice.NO,
            offer_price_eur="NaN",
        )


def test_bulgarian_price_words_are_generated_from_the_numeric_price() -> None:
    options = ContractOptions(
        contract_date=date(2026, 8, 28),
        privacy_paper_choice=BinaryChoice.NO,
        privacy_email_choice=BinaryChoice.NO,
        marketing_choice=BinaryChoice.NO,
        offer_price_eur="250 000",
    )

    assert options.offer_price_eur_words == "двеста и петдесет хиляди"
    assert bulgarian_integer_words(1_021_002) == (
        "един милион двадесет и една хиляди и две"
    )


def test_contract_options_reject_conflicting_price_words_and_decimals() -> None:
    base = {
        "contract_date": date(2026, 8, 28),
        "privacy_paper_choice": BinaryChoice.NO,
        "privacy_email_choice": BinaryChoice.NO,
        "marketing_choice": BinaryChoice.NO,
    }

    with pytest.raises(ValidationError, match="do not match"):
        ContractOptions(
            **base,
            offer_price_eur="250 000",
            offer_price_eur_words="двеста хиляди",
        )
    with pytest.raises(ValidationError, match="whole-EUR"):
        ContractOptions(**base, offer_price_eur="250000.50")


def test_generation_rejects_an_unapproved_template(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from docx import Document

    wrong_template = tmp_path / "wrong.docx"
    document = Document()
    document.add_paragraph("Valid Word file, but not an approved contract.")
    document.save(wrong_template)
    monkeypatch.setitem(TEMPLATE_BY_ROLE, ContractRole.BUYER, wrong_template)

    case_id = "2026-08-28_120500_pqr678"
    with pytest.raises(ContractGenerationError, match="approved version"):
        generate_contract(_contract_input(case_id, ContractRole.BUYER), tmp_path / case_id)


def test_contract_input_requires_identity_values() -> None:
    buyer = _contract_input("2026-08-28_120600_stu901", ContractRole.BUYER)

    with pytest.raises(ValidationError, match="Missing client values"):
        ContractInput.model_validate(
            {
                **buyer.model_dump(mode="python"),
                "client": PersonalDocument(),
            }
        )
